from __future__ import annotations

import argparse
import json
import os
import re
import tempfile
import unicodedata
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree
from pypdf import PdfReader


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
NS = {"w": W_NS, "r": R_NS}


CASES = [
    {
        "old": 1,
        "group": "term",
        "source": "Research on the scopic regime of military drones often interprets these images as products of an operative form of sensing that dehumanizes, controls, and masters the world.",
        "simulated": "对军用无人机观看机制的研究，通常把这些图像理解为一种操作性传感的产物；这种传感使世界非人化，并控制和掌握世界。",
        "final": "关于军用无人机视觉体制的研究，往往将这些图像解释为一种操作性感知形式的产物；这种感知形式会使世界去人性化，并对世界实施控制和支配。",
        "analysis": "本句的难点集中在 scopic regime、operative sensing 及三个并列动词的术语搭配。模拟初译虽能传达大意，但“观看机制”弱化了 scopic regime 所包含的视觉秩序与权力关系，“操作性传感”也与全文术语系统不合。终审译文采用“视觉体制”和“操作性感知形式”，并把 dehumanizes、controls、masters 依次处理为“去人性化、控制、支配”，使概念和批判语气都更清楚。",
    },
    {
        "old": 2,
        "group": "term",
        "source": "In contrast to the military drone, the amateur drone’s flattening can evoke interconnectedness instead of control and mastery: the drone can become a medium for planetarity.",
        "simulated": "与军用无人机相比，业余无人机的扁平化能够唤起互联，而非控制和掌控；无人机可以成为一种行星媒介。",
        "final": "与军用无人机不同，业余无人机所生成的扁平化视角能够唤起相互联结，而非控制与支配；无人机由此可以成为承载行星性的媒介。",
        "analysis": "planetarity 是本书反复使用的理论概念，不能压缩成泛指性的“行星媒介”。模拟初译还把 flattening 当作孤立名词，未说明它在此指无人机图像形成的观看效果。终审译文补出“所生成的扁平化视角”，并用“承载行星性的媒介”保留抽象概念与媒介之间的关系；interconnectedness 译为“相互联结”，也与后文的共同体论述相衔接。",
    },
    {
        "old": 3,
        "group": "term",
        "source": "This chapter argues that the drone’s flattening gaze has the potential to envision planetary communities.",
        "simulated": "本章认为，无人机的扁平化凝视有可能想象行星社区。",
        "final": "本章认为，无人机的扁平化凝视具有构想行星共同体的潜力。",
        "analysis": "planetary communities 在文中指向超越国家边界和同质化逻辑的共同体想象。“行星社区”容易被理解为具体地域或一般社会单元，不能体现其理论层级。终审译文统一采用“行星共同体”，并将 envision 译为“构想”，把原句从一般的心理想象调整为带有理论建构意味的学术判断。",
    },
    {
        "old": 4,
        "group": "term",
        "source": "By contrast, for planetary subjects, alterity is not derivative.",
        "simulated": "相比之下，对于行星主体而言，他者性并不是派生的。",
        "final": "相比之下，对于行星主体而言，他异性并非派生属性。",
        "analysis": "alterity 不是普通的“他者”指称，而是强调差异和不可同化关系的抽象概念。模拟初译采用“他者性”并非完全错误，但与终审译文中 difference、singularity 等概念的区分不够清楚；“不是派生的”也缺少明确的判断对象。终审译文将其统一为“他异性”，并补出“属性”，使 derivative 的语法作用和理论命题都更完整。",
    },
    {
        "old": 5,
        "group": "term",
        "source": "The sensorium of the drone is also more than visual: Divided presents the sounds of the ocean waves.",
        "simulated": "无人机的感官也不只是视觉：《分裂》还呈现了海浪声。",
        "final": "无人机的感知系统同样超越了视觉：《分裂》还呈现了海浪的声音。",
        "analysis": "sensorium 在此包括无人机的视觉、声音、运动和数据处理能力，并非人的单一生理“感官”。模拟初译容易把技术媒介拟人化，也没有充分表达 more than visual 的范围扩展。终审译文采用“感知系统”，并以“超越了视觉”对应后半句的声音例证，使术语含义与作品《分裂》的多模态特征相互印证。",
    },
    {
        "old": 6,
        "group": "term",
        "source": "I characterize this multiperspectival sensorium of earth drones as a form of “volumetric sensing” that grasps the three-dimensionality of space.",
        "simulated": "我把大地无人机的这种多视角感官称为一种把握空间三维性的“立体感知”。",
        "final": "我将大地无人机这种多视角的感知系统描述为一种把握空间三维性的“体积性感知”形式。",
        "analysis": "“立体感知”虽然顺口，却容易被理解为一般的双目视觉或三维成像，不能覆盖无人机在地表、空中和地下之间移动并感知空间体量的含义。终审译文以“体积性感知”对应 volumetric sensing，并延续 sensorium 的“感知系统”译法。句末补出“形式”，明确 that 从句所修饰的是一种感知方式，而不是单一技术动作。",
    },
    {
        "old": 7,
        "group": "term",
        "source": "Thus the drone functions as what media scholars have called an eco-medium, both registering environmental damage and tracing the materiality of media components.",
        "simulated": "因此，无人机发挥着媒介学者所谓“环保媒介”的作用，既记录环境破坏，也追踪媒介部件的材料。",
        "final": "因此，无人机在此充当了媒介学者们所称的“生态媒介”——它既记录了环境的破坏，又追溯了媒介组件自身的物质性。",
        "analysis": "eco-medium 讨论的是媒介与生态环境、物质基础之间的关系，并不等同于以环保为用途的工具。模拟初译中的“环保媒介”缩窄了概念范围，“媒介部件的材料”也未能表达 materiality 的理论含义。终审译文采用“生态媒介”和“媒介组件自身的物质性”，并用破折号引出双重作用，使术语解释自然进入句子。",
    },
    {
        "old": 12,
        "group": "term",
        "source": "In his analysis of planetary art, Terry Smith defines “world scale regimes of seeing” as one aesthetic configuration of the planetary.",
        "simulated": "特里·史密斯在分析行星艺术时，将“世界尺度的观看制度”定义为行星性的一种审美配置。",
        "final": "在对“行星艺术”的分析中，特里·史密斯（Terry Smith）将“世界尺度的视觉机制”界定为行星性的一种美学构型。",
        "analysis": "world scale regimes of seeing 同时涉及观看方式及其组织机制。“观看制度”过于僵硬，也容易与前文 scopic regime 的“视觉体制”混同。终审译文根据本句的概念层级采用“世界尺度的视觉机制”，并把 configuration 处理为“美学构型”；首次出现的人名补出英文，便于读者识别学者身份。",
    },
    {
        "old": 23,
        "group": "term",
        "source": "This type of visuality is found in “images for usage” (Gebrauchsbilder), such as graphs, diagrams, and maps.",
        "simulated": "这种视觉性见于图表、图解和地图等“使用图像”（Gebrauchsbilder）中。",
        "final": "这种视觉形式见于“使用型图像”（Gebrauchsbilder），如图表、图解和地图。",
        "analysis": "该句既要解释英文转述 images for usage，又要保留德语概念 Gebrauchsbilder。模拟初译中的“使用图像”像临时词组，难以显示其类别意义；列举成分置于概念之前，也削弱了术语焦点。终审译文采用“使用型图像”，保留德语原词，并把例项放在其后，使术语、来源和实例的关系一目了然。",
    },
    {
        "old": 8,
        "group": "sentence",
        "source": "These imaginaries of postcarbon communities are visions of future collectives where humans, technology, and the earth live together in a more mutually sustainable and ecological relationship.",
        "simulated": "这些后碳社区的想象构成未来集体的愿景，人类、技术和地球在其中以更可持续、更生态的关系生活。",
        "final": "这些关于后碳共同体的想象，是关于未来集体的愿景，在其中，人类、技术与地球以一种更具相互可持续性与生态性的关系共同生活。",
        "analysis": "本句的主干是 imaginaries are visions，where 从句说明愿景内部的共生关系。模拟初译把两个层次压在一个逗号句中，且省略 mutually 所强调的相互性。终审译文保留“想象—愿景”的主干，以“在其中”引出从句，并用“相互可持续性”呈现关系的双向要求；postcarbon communities 也与全文统一为“后碳共同体”。",
    },
    {
        "old": 9,
        "group": "sentence",
        "source": "These communities embody aesthetic attempts to create alternatives to the Anthropocene—or, following Bernard Stiegler, to create training grounds for the more hopeful “Neganthropocene”—providing perspectives, possibilities, and potentials for a more sustainable future and moving beyond the often dead-end and fatalistic discussions about the end of humanity.",
        "simulated": "这些共同体体现了为人类世寻找替代形式的审美尝试；或者，按照贝尔纳·斯蒂格勒的说法，它们为更有希望的“负人类世”建立训练场，为更可持续的未来提供视角、可能性和潜力，并走出关于人类终结的宿命论讨论。",
        "final": "这些共同体体现了创造“人类世替代方案”的美学尝试——或者用贝尔纳·斯蒂格勒（Bernard Stiegler）的话说，是为更具希望的“负人类世（Neganthropocene）”创造训练场——从而为更具可持续性的未来提供视角、可能性与潜力，并超越那些关于人类灭绝的死胡同式、宿命论的讨论。",
        "analysis": "原句以两个破折号嵌入理论说明，随后用两个并列的现在分词概括这些尝试产生的方向。模拟初译把插入语和主句切成松散分句，providing 与 moving beyond 的共同逻辑也不够明确。终审译文保留破折号的插入结构，用“从而”统领后半句，并补出术语英文，使论证主干、理论来源和结果关系保持连贯。",
    },
    {
        "old": 10,
        "group": "sentence",
        "source": "The planetary view does not categorize the earth according to nation-states, international cooperation, or global networks, it sees the people on the planet as one, without making them the same.",
        "simulated": "行星视角不按民族国家、国际合作或全球网络对地球进行分类；它把这个星球上的人看成一个整体，而不是同一类人。",
        "final": "行星视角不依据民族国家、国际合作或全球网络对地球进行分类，而是将地球上的人视为一个整体，但不把他们视为同一。",
        "analysis": "句末 as one 与 without making them the same 同时强调共同性和差异性。模拟初译将 same 具体化为“同一类人”，把开放的理论张力收窄成类别判断。终审译文用“而是……但……”重新组织两个层次：先说明行星视角如何看待整体，再保留“不把他们视为同一”的限制，从而避免替原文增加解释。",
    },
    {
        "old": 11,
        "group": "sentence",
        "source": "“Unmannedness” is thus a matter of relative distance and tethering, rather than of physical presence.",
        "simulated": "因此，“无人性”是相对距离和连接的问题，而不是身体是否在场的问题。",
        "final": "“无人性”因此关乎相对距离和系留方式，而不是物理在场与否。",
        "analysis": "is a matter of... rather than... 是抽象判断结构，tethering 又与前文系留气球的历史讨论直接相关。模拟初译用泛化的“连接”弱化了这一语境，两个“问题”也使句子拖沓。终审译文将主干压缩为“关乎……而不是……”，并把 tethering 明确为“系留方式”，使概念判断更紧凑。",
    },
    {
        "old": 13,
        "group": "sentence",
        "source": "Again, the drone sensorium opens up to more-than-optical modes of sensing, integrating sound, rhythm, and music.",
        "simulated": "无人机的感知系统再次进入了超光学的感知方式，并整合声音、节奏和音乐。",
        "final": "在此，无人机的感知系统再次向超越光学的感知模式敞开，将声音、节奏与音乐融合其中。",
        "analysis": "more-than-optical 不是“超光学”这一技术类别，而是指感知不再局限于光学视觉。模拟初译还把 opens up to 处理为“进入”，没有表现感知范围的展开。终审译文把复合限定语展开为“超越光学的感知模式”，用“向……敞开”承接 opens up，并将现在分词短语处理为“将……融合其中”，句内关系更自然。",
    },
    {
        "old": 14,
        "group": "sentence",
        "source": "The drone sensorium creates a volumetric interface of earth experience as it sensorially comprises the earth from multiple directions: from the ground to the sky, from within, from a hovering view, and from below.",
        "simulated": "无人机的感知系统创建了一个地球经验的体积界面，因为它从不同方向感知地球：从地面到天空、从内部、从悬停视角和从下方。",
        "final": "无人机的感知系统构成了一个体验地球的体积性界面，它能够从多个方向感知地球：从地面到天空、从内部、从悬停视角以及从下方。",
        "analysis": "句子后半 as it sensorially comprises the earth from multiple directions 如果直译为“因为它从多个方向在感官上包容了地球”，中文十分别扭，“包容”与空间感知的语义关联较弱。模拟初译的“从不同方向感知地球”反而更自然，但这说明原译存在改译效果不如初译的风险。调整后的终审译文将从句改为并列短句“它能够从多个方向感知地球”，既避免了因果连词带来的不自然感，又保留了方位序列的空间展示效果。",
    },
    {
        "old": 19,
        "group": "sentence",
        "source": "In other words, these examples of drone art demand a multiperspectival mode of interpretation where meaning oscillates without becoming fixed.",
        "simulated": "换句话说，这些无人机艺术案例需要多视角的解释方式，意义在其中不断摆动，而不会变得固定。",
        "final": "换言之，这些无人机艺术案例呼唤一种多视角的阐释模式——在这一模式中，意义不断游移，而不被固定下来。",
        "analysis": "where 从句修饰的是 mode of interpretation，而非一般地点。模拟初译虽已识别这一关系，但逗号连接使主从层次不够突出，“摆动”也偏向物理动作。终审译文以破折号引出对“阐释模式”的说明，并用“游移”对应 oscillates，使意义的不稳定状态与学术阐释语境更相称。",
    },
    {
        "old": 24,
        "group": "sentence",
        "source": "Thus my approach also relates to Paglen’s statement that drones create their own “relative geographies, folding several noncontiguous spaces around the globe into a single, distributed, battlefield.”",
        "simulated": "因此，我的方法也与帕格伦的说法有关，即无人机创造了自己的“相对地理”，把全球几个不相连的空间折叠进一个单一而分布式的战场。",
        "final": "因此，我的研究路径同样呼应了帕格伦的论断，即无人机创造了其自身的“相对地理，将全球数个互不接壤的空间折叠成了一个单一的、分布式的战场”。",
        "analysis": "该句含有作者立场、引述关系和引号内的现在分词结构。模拟初译将引号提前闭合，容易让读者误以为后半句是作者另加的解释；“方法有关”也弱化了 relates to 的论述呼应。终审译文把整段内容保留在引号内，以“研究路径同样呼应了……论断”明确引述关系，并用“互不接壤”准确表达 noncontiguous。",
    },
    {
        "old": 15,
        "group": "rhetoric",
        "source": "Seeing the world from above doesn’t just flatten things. It sharpens them",
        "simulated": "从高处看世界，不只是把事物变平，也会让事物变得锋利。",
        "final": "“从高处看世界，不仅会使其扁平化，还会使其锐化。”",
        "analysis": "flatten 与 sharpen 构成简短而突出的对照，既描述图像效果，也为下文的军事暴力批判作铺垫。模拟初译中的“变平”“变得锋利”偏具体物理状态，难以延续为媒介分析概念。终审译文使用“扁平化”和“锐化”，保留动词之间的对应关系，并用“不仅……还……”合并两句，使引语在中文中简洁有力。",
    },
    {
        "old": 16,
        "group": "rhetoric",
        "source": "As Fast’s drone pilot says, the flattened drone view sharpens things: it acts like a violent blade that cuts and kills.",
        "simulated": "正如法斯特片中的无人机飞行员所说，扁平化的无人机视角会锐化事物：它就像一把暴力的刀刃，进行切割和杀戮。",
        "final": "正如法斯特片中的无人机飞行员所言，扁平化的无人机视角会使事物变得更加尖锐：它仿佛一把暴力的利刃，切割并杀戮。",
        "analysis": "violent blade 把无人机图像的视觉“锐化”推进为暴力刀刃的隐喻。模拟初译中的“暴力的刀刃”和“进行切割和杀戮”搭配生硬，且削弱了比喻的直接性。终审译文采用“暴力的利刃”，让“切割并杀戮”直接承接刀刃意象；“会使事物变得更加尖锐”又与前句形成自然过渡。",
    },
    {
        "old": 17,
        "group": "rhetoric",
        "source": "In order to develop this argument, my first step is to historicize the drone, finding “family resemblances” between the drone and hot-air balloons.",
        "simulated": "为展开这一论点，我首先从历史角度考察无人机，寻找无人机与热气球之间的“相似关系”。",
        "final": "为展开这一论证，我首先对无人机进行历史化考察，寻找无人机与热气球之间的“家族相似性”。",
        "analysis": "family resemblances 在此不是一般的“相似关系”，而是作者组织媒介历史材料的概念隐喻。模拟初译虽然通顺，却抹去了“家族”意象及其方法论意味。终审译文保留“家族相似性”，并将 historicize 处理为“进行历史化考察”，使读者能够看出作者是在建立无人机与早期空中媒介之间的非线性关联。",
    },
    {
        "old": 18,
        "group": "rhetoric",
        "source": "Thus Drone Vision shows operative images that establish a kill grid, decide about life and death, and dehumanize the subject.",
        "simulated": "因此，《无人机视觉》展示了建立杀戮网格、决定生死并将主体非人化的操作性图像。",
        "final": "因此，《无人机视觉》展现了这样一类操作性图像：它们建立了一张“杀戮网格”，裁定生死，并对主体进行去人性化。",
        "analysis": "kill grid 既是军事目标识别结构，又带有强烈的暴力意象。模拟初译把三个动作连续置于“图像”之前，造成修饰链过长，也未突出这一核心概念。终审译文先提出“这样一类操作性图像”，再用冒号展开三项作用；“裁定生死”和“去人性化”分别对应 decide 与 dehumanize，使批判性判断更明确。",
    },
    {
        "old": 20,
        "group": "rhetoric",
        "source": "The drone is no longer an aggressive eye in the sky; it volumetrically senses alterations in the landscape due to climate change and provides a potential trajectory for sustainable landscape planning.",
        "simulated": "无人机不再是天空中富有攻击性的眼睛；它立体地感知气候变化造成的景观改变，并为可持续景观规划提供潜在轨迹。",
        "final": "无人机不再是天空中那只极具侵略性的眼睛；它以体积性感知的方式感知着由气候变化所导致的景观变迁，并为可持续的景观规划提供了一条潜在的路径。",
        "analysis": "本句把 eye in the sky 的军事隐喻与 volumetric sensing 的生态用途并置。模拟初译中的“立体地感知”和“潜在轨迹”虽接近字面，却不符合本章已经建立的术语和规划语境。终审译文保留“天空中的眼睛”这一形象，以“体积性感知的方式”延续核心术语，并把 trajectory 处理为规划意义上的“路径”，形成从军事凝视到生态实践的转向。",
    },
    {
        "old": 21,
        "group": "rhetoric",
        "source": "The artists describe this effect as a type of “terrain hugging” in which the drone seems to caress and intimately touch the earth:",
        "simulated": "艺术家把这种效果描述为一种“贴着地形飞行”，无人机仿佛在抚摸并亲密接触大地。",
        "final": "艺术家将这种效果描述为一种“贴地拥抱”，在其中，无人机似乎在爱抚并亲密地触碰大地：",
        "analysis": "terrain hugging 同时含有贴近地表的飞行状态和拥抱、触碰的身体意象。模拟初译只保留飞行方式，无法与后文 caress、touch 形成连续的触觉隐喻。终审译文采用“贴地拥抱”，随后以“爱抚并亲密地触碰大地”展开这一形象，既说明无人机的运动方式，也保留作品所强调的生态关怀。",
    },
    {
        "old": 22,
        "group": "rhetoric",
        "source": "The planet is the species of alterity, belonging to another system; and yet we inhabit it on loan.",
        "simulated": "“行星是另一个系统中的他者性物种；但我们只是暂时居住在它上面。”",
        "final": "“行星是属于另一个系统的他异性物种；然而我们只是借居于此。”",
        "analysis": "on loan 在这里不是一般的时间副词，而是非占有关系的伦理隐喻。模拟初译的“暂时居住”传达了期限，却没有表现人类对行星并无所有权的含义；“另一个系统中的”也改变了 belonging 的关系。终审译文用“借居于此”凝练地保留借用意象，并以“属于另一个系统”准确组织前半句。",
    },
    {
        "old": 25,
        "group": "rhetoric",
        "source": "Drones and their volumetric sensoria still have to be handled with care and reflection: they can integrate the human, but they can also expel it.",
        "simulated": "无人机及其体积性感知系统仍需谨慎使用和反思：它们可以把人类纳入其中，也可以把人类排除出去。",
        "final": "无人机及其体积性感知系统仍需审慎对待与反思：它们可以整合人类，但也可以排斥人类。",
        "analysis": "本句位于章节结尾，integrate 与 expel 构成对称的警示。模拟初译把 handled 限定为“使用”，使讨论范围缩小到操作层面；“把……排除出去”也偏口语。终审译文采用“审慎对待与反思”，覆盖技术应用和价值判断，再以“整合人类／排斥人类”保持两个动词的对称结构和收束力度。",
    },
]


REFERENCES = [
    "［1］Gregory, D. From a View to a Kill: Drones and Late Modern War. Theory, Culture & Society, 2011, 28(7-8): 188-215.",
    "［2］Chamayou, G. A Theory of the Drone (J. Lloyd, Trans.). New York: The New Press, 2015.",
    "［3］Agostinho, D., Maurer, K., & Veel, K. Introduction to The Sensorial Experience of the Drone. The Senses and Society, 2020, 15(3): 251-258.",
    "［4］Maurer, K. The Sensorium of the Drone and Communities. Cambridge, MA: The MIT Press, 2023.",
    "［5］Spivak, G. C. Death of a Discipline. New York: Columbia University Press, 2003.",
    "［6］Latour, B. Facing Gaia: Eight Lectures on the New Climatic Regime (C. Porter, Trans.). Cambridge: Polity Press, 2017.",
    "［7］Morton, T. Hyperobjects: Philosophy and Ecology after the End of the World. Minneapolis: University of Minnesota Press, 2013.",
    "［8］Reiss, K. Text Types, Translation Types and Translation Assessment. In A. Chesterman (Trans. & Ed.), Readings in Translation Theory. Helsinki: Finn Lectura, 1989: 105-115.",
    "［9］方梦之. 科技翻译理论的研究——十年述评与展望. 中国翻译, 1992(2): 7-10.",
    "［10］Cabré, M. T. Terminology: Theory, Methods and Applications (J. A. DeCesaris, Trans.; J. C. Sager, Ed.). Amsterdam/Philadelphia: John Benjamins, 1998.",
    "［11］Temmerman, R. Towards New Ways of Terminology Description: The Sociocognitive Approach. Amsterdam/Philadelphia: John Benjamins, 2000.",
    "［12］夏菁, 冷冰冰. 科技翻译中的术语变体及译者对策. 上海理工大学学报（社会科学版）, 2021, 43(3): 236-241.",
    "［13］杨枫, 李思伊. 什么是术语翻译谱系学？当代外语研究, 2025(5): 1-11.",
    "［14］刘亚猛. 风物常宜放眼量：西方学术文化与中西学术翻译. 中国翻译, 2004, 25(6): 44-48.",
    "［15］连淑能. 英汉对比研究：增订本. 北京: 高等教育出版社, 2010.",
    "［16］Murtisari, E. T. Explicitation in Translation Studies: The Journey of an Elusive Concept. Translation & Interpreting, 2016, 8(2): 64-81.",
    "［17］柯飞. 翻译中的隐和显. 外语教学与研究, 2005, 37(4): 303-307.",
    "［18］Klaudy, K. Explicitation. In M. Baker & G. Saldanha (Eds.), Routledge Encyclopedia of Translation Studies (2nd ed.). London/New York: Routledge, 2009: 104-109.",
    "［19］Lakoff, G., & Johnson, M. Metaphors We Live By. Chicago: University of Chicago Press, 1980.",
    "［20］Hatim, B., & Mason, I. Discourse and the Translator. London: Longman, 1990.",
    "［21］Appiah, K. A. Thick Translation. Callaloo, 1993, 16(4): 808-819.",
    "［22］国家市场监督管理总局, 国家标准化管理委员会. 翻译服务 第1部分：笔译服务要求: GB/T 19363.1—2022[S]. 北京: 中国标准出版社, 2022.",
    "［23］Farocki, H. Phantom Images. Public, 2004(29): 12-22.",
]


TOC_TITLES = [
    ("摘 要", 1, "front"),
    ("ABSTRACT", 1, "front"),
    ("第一章 引言", 1, "body"),
    ("1.1 研究背景及意义", 2, "body"),
    ("1.2 研究问题", 2, "body"),
    ("1.3 报告结构", 2, "body"),
    ("第二章 《无人机感知与共同体》翻译项目概述", 1, "body"),
    ("2.1 项目简介", 2, "body"),
    ("2.2 翻译流程", 2, "body"),
    ("2.2.1 译前准备", 3, "body"),
    ("2.2.2 翻译过程", 3, "body"),
    ("2.2.3 译后管理", 3, "body"),
    ("第三章 《无人机感知与共同体》翻译项目案例分析", 1, "body"),
    ("3.1 源语文本的类型与特征", 2, "body"),
    ("3.1.1 词汇特征：跨学科概念密集", 3, "body"),
    ("3.1.2 句法特征：复合结构与隐含关系", 3, "body"),
    ("3.1.3 语篇特征：隐喻、互文与批判性表达", 3, "body"),
    ("3.2 翻译中的主要问题", 2, "body"),
    ("3.2.1 跨学科术语与理论表达", 3, "body"),
    ("3.2.2 复杂句与论证关系", 3, "body"),
    ("3.2.3 隐喻、修辞与语篇表达", 3, "body"),
    ("3.3 具体处理", 2, "body"),
    ("3.3.1 术语查证与译名统一", 3, "body"),
    ("3.3.2 句子重组与逻辑梳理", 3, "body"),
    ("3.3.3 隐喻保留与语篇调整", 3, "body"),
    ("3.4 小结", 2, "body"),
    ("第四章 总结与反思", 1, "body"),
    ("参考文献", 1, "body"),
    ("致 谢", 1, "body"),
    ("附录一 《无人机感知与共同体》原文与译文", 1, "body"),
]


def norm_text(text: str) -> str:
    # Remove Word footnote markers before NFKC turns them into ordinary digits.
    text = re.sub(r"[¹²³⁴⁵⁶⁷⁸⁹⁰\u2060]", "", text)
    text = unicodedata.normalize("NFKC", text)
    # Word's macOS PDF exporter uses this CJK-radical glyph for 目 without
    # a Unicode compatibility decomposition.
    text = text.replace("⺫", "目")
    text = text.replace("’", "'").replace("‘", "'").replace("“", '"').replace("”", '"')
    text = re.sub(r"\s+", "", text)
    return text.strip('"')


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_east_asia(run, east: str, latin: str = "Times New Roman", size: float | None = None) -> None:
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:cs"), latin)


def set_style_fonts(style, east: str, latin: str, size: float, bold: bool | None = None) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:cs"), latin)


def set_outline_level(style, value: int | None) -> None:
    ppr = style.element.get_or_add_pPr()
    node = ppr.find(qn("w:outlineLvl"))
    if value is None:
        if node is not None:
            ppr.remove(node)
        return
    if node is None:
        node = OxmlElement("w:outlineLvl")
        ppr.append(node)
    node.set(qn("w:val"), str(value))


def get_or_add_style(doc, name: str, base: str = "Normal"):
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base]
        return style


def set_fixed_20(paragraph, before: float = 0, after: float = 0) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def configure_styles(doc) -> None:
    normal = doc.styles["Normal"]
    set_style_fonts(normal, "Songti SC", "Times New Roman", 10.5, False)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(20)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    h1 = doc.styles["Heading 1"]
    set_style_fonts(h1, "Heiti SC", "Times New Roman", 15, True)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.first_line_indent = Pt(0)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h1.paragraph_format.line_spacing = Pt(20)
    h1.paragraph_format.space_before = Pt(30)
    h1.paragraph_format.space_after = Pt(30)
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.keep_with_next = True
    set_outline_level(h1, 0)

    h2 = doc.styles["Heading 2"]
    set_style_fonts(h2, "Heiti SC", "Times New Roman", 14, True)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.first_line_indent = Pt(0)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h2.paragraph_format.line_spacing = Pt(20)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(10)
    h2.paragraph_format.keep_with_next = True
    set_outline_level(h2, 1)

    h3 = doc.styles["Heading 3"]
    set_style_fonts(h3, "Heiti SC", "Times New Roman", 12, True)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.first_line_indent = Pt(0)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h3.paragraph_format.line_spacing = Pt(20)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(10)
    h3.paragraph_format.keep_with_next = True
    set_outline_level(h3, 2)

    h4 = doc.styles["Heading 4"]
    set_style_fonts(h4, "Heiti SC", "Times New Roman", 12, True)
    h4.paragraph_format.first_line_indent = Pt(0)
    h4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h4.paragraph_format.line_spacing = Pt(20)
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(10)
    h4.paragraph_format.keep_with_next = True
    set_outline_level(h4, 3)

    front = get_or_add_style(doc, "FrontTitle")
    front.base_style = doc.styles["Normal"]
    set_style_fonts(front, "Heiti SC", "Times New Roman", 14, True)
    front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.paragraph_format.first_line_indent = Pt(0)
    front.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    front.paragraph_format.line_spacing = Pt(20)
    front.paragraph_format.space_before = Pt(0)
    front.paragraph_format.space_after = Pt(20)
    front.paragraph_format.keep_with_next = True
    set_outline_level(front, 8)

    front_toc = get_or_add_style(doc, "FrontTocTitle")
    front_toc.base_style = front
    set_style_fonts(front_toc, "Heiti SC", "Times New Roman", 14, True)
    front_toc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front_toc.paragraph_format.first_line_indent = Pt(0)
    front_toc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    front_toc.paragraph_format.line_spacing = Pt(20)
    front_toc.paragraph_format.space_before = Pt(0)
    front_toc.paragraph_format.space_after = Pt(20)
    front_toc.paragraph_format.keep_with_next = True
    set_outline_level(front_toc, 8)

    case_label = get_or_add_style(doc, "CaseLabel")
    case_label.base_style = doc.styles["Normal"]
    set_style_fonts(case_label, "Heiti SC", "Times New Roman", 10.5, True)
    case_label.paragraph_format.first_line_indent = Pt(0)
    case_label.paragraph_format.space_before = Pt(10)
    case_label.paragraph_format.space_after = Pt(0)
    case_label.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    case_label.paragraph_format.line_spacing = Pt(20)
    case_label.paragraph_format.keep_with_next = True
    set_outline_level(case_label, None)

    case_text = get_or_add_style(doc, "CaseText")
    case_text.base_style = doc.styles["Normal"]
    set_style_fonts(case_text, "Songti SC", "Times New Roman", 10.5, False)
    case_text.paragraph_format.first_line_indent = Pt(0)
    case_text.paragraph_format.space_before = Pt(0)
    case_text.paragraph_format.space_after = Pt(0)
    case_text.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    case_text.paragraph_format.line_spacing = Pt(20)
    set_outline_level(case_text, None)

    caption = get_or_add_style(doc, "Caption")
    caption.base_style = doc.styles["Normal"]
    set_style_fonts(caption, "Songti SC", "Times New Roman", 10.5, False)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    caption.paragraph_format.line_spacing = Pt(20)
    caption.paragraph_format.space_before = Pt(10)
    caption.paragraph_format.space_after = Pt(0)
    caption.paragraph_format.keep_with_next = True
    set_outline_level(caption, None)

    refs = get_or_add_style(doc, "References")
    # The source document's custom References style was based on Heading 1,
    # which made every bibliography entry a TOC item and a new page in Word.
    refs.base_style = doc.styles["Normal"]
    set_style_fonts(refs, "Songti SC", "Times New Roman", 10.5, False)
    refs.paragraph_format.first_line_indent = Pt(-21)
    refs.paragraph_format.left_indent = Pt(21)
    refs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    refs.paragraph_format.line_spacing = Pt(20)
    refs.paragraph_format.space_before = Pt(0)
    refs.paragraph_format.space_after = Pt(0)
    set_outline_level(refs, None)

    for level in (1, 2, 3):
        name = f"toc {level}"
        toc = get_or_add_style(doc, name)
        toc.base_style = doc.styles["Normal"]
        set_style_fonts(toc, "Songti SC", "Times New Roman", 10.5, False)
        toc.paragraph_format.left_indent = Cm(0.74 * (level - 1))
        toc.paragraph_format.first_line_indent = Pt(0)
        toc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        toc.paragraph_format.line_spacing = Pt(20)
        toc.paragraph_format.space_before = Pt(0)
        toc.paragraph_format.space_after = Pt(0)
        set_outline_level(toc, None)

    tof = get_or_add_style(doc, "TableListEntry")
    tof.base_style = doc.styles["Normal"]
    set_style_fonts(tof, "Songti SC", "Times New Roman", 10.5, False)
    tof.paragraph_format.first_line_indent = Pt(0)
    tof.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    tof.paragraph_format.line_spacing = Pt(20)
    tof.paragraph_format.space_before = Pt(0)
    tof.paragraph_format.space_after = Pt(0)
    set_outline_level(tof, None)


def set_text(paragraph, text: str, style: str | None = None) -> None:
    clear_paragraph(paragraph)
    if style:
        paragraph.style = style
    run = paragraph.add_run(text)
    return run


def set_text_with_citations(paragraph, text: str, style: str | None = None) -> None:
    clear_paragraph(paragraph)
    if style:
        paragraph.style = style
    for token in re.split(r"(\[\[CITE:\d+\]\])", text):
        m = re.fullmatch(r"\[\[CITE:(\d+)\]\]", token)
        if m:
            run = paragraph.add_run(f"［{m.group(1)}］")
            run.font.superscript = True
        elif token:
            paragraph.add_run(token)


def find_paragraph(doc, exact: str | None = None, starts: str | None = None):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if exact is not None and re.sub(r"\s+", "", text) == re.sub(r"\s+", "", exact):
            return paragraph
        if starts is not None and text.startswith(starts):
            return paragraph
    raise KeyError(exact or starts)


def set_paragraph_text(doc, starts: str, text: str, style: str | None = None, cites: bool = False):
    matches = [
        p
        for p in doc.paragraphs
        if not p.style.name.lower().startswith("toc") and p.text.strip().startswith(starts)
    ]
    if not matches:
        raise KeyError(starts)
    p = matches[0]
    if cites:
        set_text_with_citations(p, text, style)
    else:
        set_text(p, text, style)
    return p


def set_paragraph_text_nth(doc, starts: str, occurrence: int, text: str, style: str | None = None, cites: bool = False):
    matches = [
        p
        for p in doc.paragraphs
        if not p.style.name.lower().startswith("toc") and p.text.strip().startswith(starts)
    ]
    if occurrence >= len(matches):
        raise KeyError(f"Missing occurrence {occurrence} for {starts!r}")
    p = matches[occurrence]
    if cites:
        set_text_with_citations(p, text, style)
    else:
        set_text(p, text, style)
    return p


def delete_range_between(start_element, end_element) -> None:
    parent = start_element.getparent()
    current = start_element
    while current is not None and current is not end_element:
        nxt = current.getnext()
        parent.remove(current)
        current = nxt


def move_before(anchor, element) -> None:
    anchor.addprevious(element)


def add_paragraph_before(doc, anchor, text: str = "", style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        paragraph.add_run(text)
    move_before(anchor, paragraph._p)
    return paragraph


def add_cited_paragraph_before(doc, anchor, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    set_text_with_citations(paragraph, text, style)
    move_before(anchor, paragraph._p)
    return paragraph


def append_simple_field(paragraph, instruction: str, cached: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = cached
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_case(doc, anchor, number: int, case: dict) -> None:
    label = add_paragraph_before(doc, anchor, f"例[{number}]", "CaseLabel")
    label.paragraph_format.keep_with_next = True
    for idx, (prefix, value) in enumerate(
        [("原文：", case["source"]), ("模拟初译：", case["simulated"]), ("改译：", case["final"]), ("分析：", case["analysis"])]
    ):
        p = add_paragraph_before(doc, anchor, style="CaseText")
        r0 = p.add_run(prefix)
        r0.bold = True
        p.add_run(value)
        if idx < 3:
            p.paragraph_format.keep_with_next = True
        if prefix == "原文：":
            for run in p.runs[1:]:
                set_east_asia(run, "Songti SC", "Times New Roman", 10.5)
        if prefix == "分析：":
            p.paragraph_format.space_after = Pt(10)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def build_feature_table(doc, anchor) -> None:
    caption = add_paragraph_before(doc, anchor, style="Caption")
    caption.add_run("表3.")
    append_simple_field(caption, "SEQ 表 \\* ARABIC", "1")
    caption.add_run(" 源语文本的主要特征及其对翻译实践的影响")

    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths_cm = [3.7, 5.5, 6.0]
    data = [
        ["文本特征", "具体表现", "对翻译的影响"],
        ["跨学科术语密集", "媒介研究、生态哲学与无人机技术概念交织", "需要查证概念来源并统一译名"],
        ["复合句结构较多", "从句、插入语和抽象名词结构承载多层关系", "需要重组语序并梳理论证关系"],
        ["隐喻与互文突出", "视觉、暴力、触觉隐喻及跨媒介引用频繁", "需要判断修辞功能和必要的背景信息"],
        ["批判性立场鲜明", "军事凝视、资本逻辑与技术乐观主义反复出现", "需要准确处理转折、评价和语篇推进"],
    ]
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = data[r_idx][c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            twips = int(Cm(widths_cm[c_idx]).twips)
            set_cell_width(cell, twips)
            p = cell.paragraphs[0]
            p.style = doc.styles["Normal"]
            p.paragraph_format.first_line_indent = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_fixed_20(p)
            for run in p.runs:
                set_east_asia(run, "Songti SC", "Times New Roman", 10.5)
                run.bold = r_idx == 0
            if r_idx == 0:
                set_cell_shading(cell, "E7E6E6")
    set_repeat_header(table.rows[0])
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(Cm(sum(widths_cm)).twips)))
    tbl_w.set(qn("w:type"), "dxa")
    move_before(anchor, table._tbl)


def rebuild_chapter_three(doc) -> None:
    start = find_paragraph(doc, exact="第三章 《无人机感知与共同体》翻译项目案例分析")
    chapter4 = find_paragraph(doc, exact="第四章 总结与反思")
    anchor = chapter4._p
    delete_range_between(start._p, anchor)

    add_paragraph_before(doc, anchor, "第三章 《无人机感知与共同体》翻译项目案例分析", "Heading 1")
    add_cited_paragraph_before(
        doc,
        anchor,
        "本章以《无人机感知与共同体》第三部分的英汉翻译实践为对象。依据莱斯关于文本功能的讨论[[CITE:8]]，源文本可视为兼具信息功能与表达功能的学术文本：它既传递无人机感知、遥感技术和生态媒介等知识，也借助隐喻、互文和艺术案例展开技术批判。下文先说明源文本的主要特征，再把翻译问题归纳为术语与理论表达、复杂句与论证关系、隐喻修辞与语篇表达三类，最后结合二十五个例证讨论具体处理。",
        "Normal",
    )

    add_paragraph_before(doc, anchor, "3.1 源语文本的类型与特征", "Heading 2")
    add_cited_paragraph_before(
        doc,
        anchor,
        "源文本选自凯瑟琳·毛雷尔的学术专著《无人机感知与共同体》[[CITE:4]]。该部分以无人机、热气球、遥感技术和实验艺术为材料，讨论技术感知如何参与人类对地球及共同体的理解。文本并非单纯说明技术原理，而是在媒介研究、环境人文学与技术哲学之间展开论证，因此译文既要准确传递知识，也要保留作者的判断和表达方式。",
        "Normal",
    )
    add_paragraph_before(doc, anchor, "3.1.1 词汇特征：跨学科概念密集", "Heading 3")
    add_cited_paragraph_before(
        doc,
        anchor,
        "源文本中的 planetarity、sensorium、volumetric sensing、alterity、operative visuality 等词语来自不同学科，且在本书中彼此关联。术语学研究指出，术语是专业知识结构中的概念单位，其具体含义会受到使用语境影响[[CITE:10]][[CITE:11]]。因此，译者不能只按词典义逐词对应，而要查明概念来源、辨别相近术语，并在正文、案例和术语表之间保持一致。",
        "Normal",
    )
    add_paragraph_before(doc, anchor, "3.1.2 句法特征：复合结构与隐含关系", "Heading 3")
    add_cited_paragraph_before(
        doc,
        anchor,
        "源文本经常借助定语从句、非谓语结构、插入语和介词短语，把主张、限定与补充说明压缩在一个句子中。英汉语言在连接方式和信息组织上存在差异[[CITE:15]]；翻译时需要先辨认句子主干，再判断因果、转折、让步或递进关系是否需要在中文中明示。显性化并非任意增添解释，而是在不改变原意的前提下，使原文已有的关系在译文中可辨[[CITE:16]][[CITE:17]][[CITE:18]]。",
        "Normal",
    )
    add_paragraph_before(doc, anchor, "3.1.3 语篇特征：隐喻、互文与批判性表达", "Heading 3")
    add_cited_paragraph_before(
        doc,
        anchor,
        "源文本中的 flattening、violent blade、terrain hugging、on loan 等表达不仅提供形象，也参与论证。概念隐喻研究表明，隐喻能够组织对抽象经验的理解[[CITE:19]]；语篇分析则要求译者结合上下文判断表达的立场和功能[[CITE:20]]。此外，Gebrauchsbilder、relative geographies 等词语还带有摄影史或媒介理论背景，必要时需要保留原词或在句内补足最少的说明[[CITE:21]]。",
        "Normal",
    )
    build_feature_table(doc, anchor)

    add_paragraph_before(doc, anchor, "3.2 翻译中的主要问题", "Heading 2")
    add_paragraph_before(doc, anchor, "为使案例分类与实际翻译决策对应，本节将原有五类问题整合为三个边界较清楚的方面。三类问题并非互相隔绝：一个长句可能同时包含术语或隐喻，但分类时以该例中最需要说明的处理为准。", "Normal")
    add_paragraph_before(doc, anchor, "3.2.1 跨学科术语与理论表达", "Heading 3")
    add_paragraph_before(doc, anchor, "第一类问题是核心术语和理论表达的确定。planetarity 不能与 globalism 混同；sensorium 不能直接等同于人的“感官”；volumetric sensing 也不只是普通“立体视觉”。类似地，scopic regime、operative visuality、alterity、eco-medium 和 Gebrauchsbilder 都需要结合本书的论证位置判断。处理此类词语时，主要困难在于概念辨析、术语间的边界以及全文译名的一致。", "Normal")
    add_paragraph_before(doc, anchor, "3.2.2 复杂句与论证关系", "Heading 3")
    add_paragraph_before(doc, anchor, "第二类问题来自复合句及其中的论证关系。源文本常在一个句子中并置主张、插入说明、例证和评价，或用 where、as、rather than、现在分词等结构压缩关系。若完全沿用英语语序，中文容易出现修饰范围不清、层次松散或重心偏移；若解释过多，又会把原文开放的判断改写成确定结论。因此，句子重组必须以源文已有关系为边界。", "Normal")
    add_paragraph_before(doc, anchor, "3.2.3 隐喻、修辞与语篇表达", "Heading 3")
    add_paragraph_before(doc, anchor, "第三类问题涉及隐喻、互文和批判性语篇。军事无人机被写成“天空中的眼睛”“暴力的利刃”和“杀戮网格”，生态艺术中的无人机则会“贴地拥抱”并触碰大地。这些形象承担不同的评价功能，不能一律直译或一律解释。译者还需准确处理转折和对照，避免因追求顺畅而削弱作者对军事凝视、技术乐观主义和生态剥削的批判。", "Normal")

    add_paragraph_before(doc, anchor, "3.3 具体处理", "Heading 2")
    add_paragraph_before(doc, anchor, "以下二十五个例证均取自源文本，例中“改译”与终审译文保持一致。为便于展示翻译决策，“模拟初译”依据源文表层结构和人类译者在术语查证不足、句法关系识别不充分或修辞功能判断不清时可能出现的常见处理构拟，仅用于分析对照，不代表真实历史稿。", "Normal")

    groups = [
        ("term", "3.3.1 术语查证与译名统一", "本节九个案例分别涉及视觉研究、行星性、共同体、他异性、感知系统、体积性感知、生态媒介和跨语际术语。处理重点是先确认概念在本书中的指向，再选择能够进入中文学术语篇的译名，并检查相关术语之间是否前后一致。"),
        ("sentence", "3.3.2 句子重组与逻辑梳理", "本节八个案例主要考察关系从句、对比结构、插入语、并列非谓语和引述结构。分析不以“拆句”本身为目的，而是说明译文如何重新安排信息层次、保持论证关系，并控制明示的范围。"),
        ("rhetoric", "3.3.3 隐喻保留与语篇调整", "本节八个案例涉及视觉、暴力、家族、触觉和借居等隐喻，以及章节结尾的批判性对照。处理时既要保留原文可识别的形象，也要根据中文搭配和语篇位置调整表达，使修辞仍能服务于论证。"),
    ]
    number = 1
    for key, heading, intro in groups:
        add_paragraph_before(doc, anchor, heading, "Heading 3")
        add_paragraph_before(doc, anchor, intro, "Normal")
        for case in [c for c in CASES if c["group"] == key]:
            add_case(doc, anchor, number, case)
            number += 1
        closing = {
            "term": "上述案例表明，术语处理需要同时考虑来源、语境和词项之间的关系。译名一经确定，还应在摘要、正文、案例和术语材料中复核，避免同一概念出现无根据的变体。",
            "sentence": "这些案例显示，复杂句处理的关键是先还原原句的主干和关系，再决定语序、连接方式和标点。必要的明示可以帮助组织中文句子，但不能替原文补写结论。",
            "rhetoric": "这些案例说明，修辞处理不能脱离论证位置。军事语境中的形象需要保留警示和暴力色彩，生态艺术语境中的触觉表达则要保留亲近大地的意味；章节结尾的对照还应保持克制而明确的批判力度。",
        }[key]
        add_paragraph_before(doc, anchor, closing, "Normal")

    add_paragraph_before(doc, anchor, "3.4 小结", "Heading 2")
    add_paragraph_before(doc, anchor, "本章根据源文本的实际特征，将翻译问题整合为跨学科术语与理论表达、复杂句与论证关系、隐喻修辞与语篇表达三类，并以二十五个案例说明具体处理。三个案例组分别包含九例、八例和八例，均以终审译文为准。", "Normal")
    add_paragraph_before(doc, anchor, "案例表明：术语译名应结合概念来源、上下文和全文用法确定；复杂句应先理清主干与论证关系，再调整语序、连接方式和标点；隐喻与语篇表达则应根据评价功能和段落位置处理。上述做法均以终审译文和具体语境为依据，并非可以机械套用的策略清单。", "Normal")


def update_front_and_body(doc) -> None:
    # Covers and declarations
    set_paragraph_text(doc, "二О二X年X月", "【待作者填写】")
    set_paragraph_text(doc, "February 2026", "【待作者填写】")
    set_paragraph_text(doc, "研究生签名：", "研究生签名：【待作者填写】    日期：【待作者填写】")
    signature_lines = [p for p in doc.paragraphs if p.text.strip().startswith("研究生签名：")]
    if len(signature_lines) > 1:
        set_text(signature_lines[1], "研究生签名：【待作者填写】    导师签名：【待作者填写】    日期：【待作者填写】")
    if doc.tables and len(doc.tables[0].rows) >= 5:
        doc.tables[0].cell(4, 1).text = "【待作者填写】"

    # Abstracts
    set_paragraph_text(
        doc,
        "本翻译实践报告以丹麦学者",
        "本翻译实践报告以丹麦学者凯瑟琳·毛雷尔（Kathrin Maurer）学术专著 The Sensorium of the Drone and Communities 第三部分“The Earth”的英汉翻译实践为研究对象。源文本处于媒介研究、环境人文学与技术哲学的交叉领域，主要探讨无人机感知如何参与行星共同体、后碳共同体以及人类世生态批判的建构。该文本兼具信息型与表达型特征，既包含无人机遥感、景观测绘和多光谱成像等技术知识，也通过艺术案例、哲学概念和空间隐喻表达作者对军事凝视、技术乐观主义和生态危机的批判。",
    )
    set_paragraph_text(
        doc,
        "基于源文本的跨学科属性",
        "结合翻译过程与导师意见，本报告将主要问题整合为三类：跨学科术语与理论表达的处理、复杂句与论证关系的处理，以及隐喻、修辞和相关语篇表达的处理。报告以二十五个真实源文例句为依据，通过术语查证与译名统一、句子重组与逻辑梳理、隐喻保留与语篇调整等具体方法，说明译文如何在准确传递知识的同时保持原文的论证层次和批判立场。",
    )
    set_paragraph_text(
        doc,
        "通过本次翻译实践",
        "本次实践表明，跨学科学术文本翻译需要把术语、句法和语篇放在同一论证环境中判断。译者既要核对概念及专名，也要控制解释的范围，并通过多轮双语复核保证案例分析与最终译文一致。本报告希望为媒介研究、环境人文学和技术哲学类英文学术文本的汉译提供具体参考。",
    )
    set_paragraph_text(doc, "关键词：", "关键词：《无人机感知与共同体》，学术文本翻译，术语统一，句子重组，修辞表达")

    set_paragraph_text(
        doc,
        "This translation practice report is based",
        "This translation practice report examines the English-to-Chinese translation of Part III, “The Earth,” from Kathrin Maurer’s academic monograph The Sensorium of the Drone and Communities. At the intersection of media studies, environmental humanities, and philosophy of technology, the source text combines technical discussion of drone sensing and mapping with critiques of military vision, techno-optimism, and ecological crisis.",
    )
    set_paragraph_text(
        doc,
        "Given the interdisciplinary nature",
        "The report groups the main translation problems into three categories: interdisciplinary terminology and theoretical expressions, complex sentences and argumentative relations, and metaphorical, rhetorical, and related discourse expressions. Twenty-five authentic source-text examples are used to explain terminology verification and consistency, sentence restructuring and logical clarification, and the adjustment of metaphor and discourse. These procedures support accurate knowledge transfer while preserving the source text’s argumentative structure and critical stance.",
    )
    set_paragraph_text(
        doc,
        "This translation practice shows",
        "The practice shows that terminological, syntactic, and discourse decisions must be made within the same argumentative context. Concepts and proper names require verification, explanatory additions should remain limited, and repeated bilingual review is necessary to keep the case analyses consistent with the final translation. The report offers a concrete reference for translating English academic texts in related interdisciplinary fields.",
    )
    set_paragraph_text(doc, "Keywords:", "Keywords: The Sensorium of the Drone and Communities, academic text translation, terminology consistency, sentence restructuring, rhetorical expression")

    # Chapter 1 headings and aligned research questions
    set_paragraph_text(doc, "第一章 引言", "第一章 引言", "Heading 1")
    set_paragraph_text(doc, "1.1 研究背景及意义", "1.1 研究背景及意义", "Heading 2")
    set_paragraph_text(doc, "1.2 研究问题", "1.2 研究问题", "Heading 2")
    set_paragraph_text(doc, "1.3 报告结构", "1.3 报告结构", "Heading 2")

    set_paragraph_text(
        doc,
        "近年来，无人机技术",
        "近年来，无人机技术在军事侦察、遥感测绘、环境监测、影像生产和艺术创作等领域持续扩展，其社会影响已超出单纯的工程技术范畴。Gregory 在讨论无人机与晚期现代战争时指出，无人机是理解“远程观看”与“远程杀伤”关系的重要技术对象[[CITE:1]]；Chamayou 也从哲学与政治伦理角度分析了无人机如何改变战争主体、战场空间和杀伤责任的分配方式[[CITE:2]]。随着媒介研究、环境人文学和后人类主义理论的发展，研究者开始进一步关注无人机如何通过视觉、声音、数据处理和空间测绘参与人类对环境的感知。Agostinho、Maurer 与 Veel 据此提出，应超越视觉中心主义，将无人机理解为参与更广泛感知组合的媒介技术[[CITE:3]]。由此可见，无人机不只是军事工具或遥感设备，也逐渐成为连接技术、空间与生态问题的重要研究对象。",
        cites=True,
    )
    set_paragraph_text(
        doc,
        "本翻译实践项目选取丹麦学者凯瑟琳·毛雷尔",
        "本翻译实践项目选取丹麦学者凯瑟琳·毛雷尔（Kathrin Maurer）的学术专著 The Sensorium of the Drone and Communities 第三部分“The Earth”为源文本[[CITE:4]]。该书由麻省理工学院出版社于 2023 年出版，主要讨论无人机感知系统与共同体想象之间的关系。第三部分“The Earth”围绕无人机对地球的观看、测绘和感知展开，重点论述“扁平化感知”和“体积性感知”如何关联行星共同体、后碳共同体和人类世生态批判。该部分既包含无人机遥感、摄影测量、激光雷达和点云模型等技术知识，也涉及斯皮瓦克的“行星性”[[CITE:5]]、拉图尔的“盖娅”[[CITE:6]]和莫顿的“超客体”[[CITE:7]]等理论资源。笔者承担该部分文本的英汉翻译、术语整理和译后修订工作，源文本约 17,000 词。",
        cites=True,
    )
    set_paragraph_text(
        doc,
        "从文本性质看",
        "从文本性质看，该项目并非一般科技说明文本，也不是单纯的人文学术论述。依据莱斯文本类型理论，翻译时应关注源文本的主导功能及其在目标语中的实现方式[[CITE:8]]。本项目源文本一方面具有明显的信息功能，需要准确传达无人机技术和遥感实践相关知识；另一方面又具有较强的表达功能和论辩功能，作者通过隐喻、互文和艺术案例表达对军事凝视、技术乐观主义和生态危机的反思。因此，译者面对的不是单纯的技术信息转写，而是技术术语、理论概念和批判性修辞相互交织的学术话语。",
        cites=True,
    )
    set_paragraph_text(
        doc,
        "从国内翻译实践背景看",
        "从国内翻译实践背景看，科技翻译通常强调信息准确、术语规范和表达简明[[CITE:9]]。然而，本项目中的无人机技术知识并非以说明书、技术标准或工程报告的形式出现，而是嵌入媒介理论、环境人文学和艺术批评的论证之中。这使译者不能只依赖既有科技术语库，也不能完全按照一般人文学术文本处理其修辞表达，而需要在技术准确性、理论来源和中文学术表达之间建立平衡。尤其是 planetarity、sensorium、volumetric sensing、necroethics、postcarbon communities 等概念，在中文语境中未必存在稳定译名，若仅按字面意义处理，容易造成概念误读或理论内涵弱化。",
        cites=True,
    )
    set_paragraph_text(
        doc,
        "基于上述背景，本报告围绕",
        "基于上述背景，本报告围绕《无人机感知与共同体》第三部分“The Earth”的英汉翻译实践展开，重点讨论三类问题：跨学科术语与理论表达、复杂句与论证关系，以及隐喻、修辞和相关语篇表达。通过对二十五个案例的分析，本文旨在总结本项目中有明确证据的处理方法，为同类科技人文交叉文本的英汉翻译提供具体参考。",
    )
    set_paragraph_text(
        doc,
        "本报告基于《无人机感知与共同体》第三部分",
        "本报告基于《无人机感知与共同体》第三部分“The Earth”的英汉翻译实践。源文本约 17,000 词，内容处于媒介研究、环境人文学与技术哲学的交叉地带。根据翻译过程中反复出现的问题，并结合第三章的案例分类，本报告提出以下三个研究问题：",
    )
    set_paragraph_text(doc, "（1）面对跨学科", "（1）面对跨学科术语与理论表达，译者如何确定译名并保持概念一致？")
    set_paragraph_text(
        doc,
        "源文本中包含较多具有理论谱系",
        "源文本包含 planetarity、sensorium、volumetric sensing、alterity、postcarbon communities 等跨学科概念。Cabré 将术语视为专业知识结构中的概念单位[[CITE:10]]，Temmerman 也强调术语意义具有语境性和动态性[[CITE:11]]；相关研究还讨论了科技翻译中的术语变体[[CITE:12]]、术语翻译的知识迁移[[CITE:13]]和西方学术概念的复杂性[[CITE:14]]。因此，本研究关注译者如何查明概念来源、区分相近词项并统一全文译名。",
        cites=True,
    )
    set_paragraph_text(doc, "（2）面对英语学术长句", "（2）面对复杂句和隐含的论证关系，译者如何重组句子而不过度解释？")
    set_paragraph_text(
        doc,
        "源文本中大量句子具有明显的英语形合特征",
        "源文本常通过定语从句、非谓语结构、插入语和介词短语压缩多层信息。英汉语言在连接方式和信息组织上存在差异[[CITE:15]]，显性化也会受到语言结构、文体和语用等因素影响[[CITE:16]][[CITE:17]][[CITE:18]]。因此，本研究关注译者如何识别句子主干和从属关系，通过语序、标点或连接语梳理论证，同时避免添加原文没有的判断。",
        cites=True,
    )
    set_paragraph_text(doc, "（3）面对源文本中的隐喻表达", "（3）面对隐喻、互文和批判性表达，译者如何在中文语篇中保留其功能？")
    set_paragraph_text(
        doc,
        "源文本中的隐喻并非单纯的文学装饰",
        "源文本中的 violent blade、eye in the sky、terrain hugging 和 on loan 等表达既提供形象，也参与作者的技术批判和生态论证。概念隐喻研究指出，隐喻会参与抽象经验的组织[[CITE:19]]；语篇分析强调语言选择与语境、立场之间的关系[[CITE:20]]，必要的背景说明也有助于目标读者识别互文信息[[CITE:21]]。本研究据此关注译者如何保留关键形象、处理理论来源，并通过转折、对照和评价词再现原文的批判性立场。",
        cites=True,
    )
    # Remove obsolete question paragraphs 4 and 5.
    for prefix in ("（4）面对互文", "源文本中包含较多跨媒介", "（5）面对源文本", "源文本并非中立介绍"):
        try:
            remove_paragraph(find_paragraph(doc, starts=prefix))
        except KeyError:
            pass

    set_paragraph_text(doc, "第一章为引言", "第一章为引言，主要介绍本次翻译实践的研究背景、研究意义、三个研究问题和报告结构。")
    set_paragraph_text(doc, "第三章为翻译项目案例分析", "第三章为翻译项目案例分析，是本报告的核心部分。本章先说明源语文本的类型与特征，再将主要问题整合为跨学科术语与理论表达、复杂句与论证关系、隐喻修辞与语篇表达三类，并以二十五个案例讨论相应的具体处理。")
    set_paragraph_text(doc, "第四章为总结与反思", "第四章为总结与反思，主要概括本次实践的发现，回应引言提出的三个研究问题，并说明术语查证、句子重组、修辞处理和译后复核方面仍存在的不足。")
    set_paragraph_text(doc, "参考文献部分列出", "参考文献部分列出本报告实际引用的文献。附录一保留原文与译文对照材料的标题结构。")

    # Chapter 2 local, evidence-based cleanup.
    set_paragraph_text_nth(
        doc,
        "本翻译实践项目选取丹麦学者凯瑟琳·毛雷尔",
        1,
        "本翻译实践项目选取丹麦学者凯瑟琳·毛雷尔（Kathrin Maurer）的学术专著 The Sensorium of the Drone and Communities 第三部分“The Earth”为源文本[[CITE:4]]。该书由麻省理工学院出版社于 2023 年出版，主要探讨无人机感知系统与共同体想象之间的关系。截至本报告撰写时，笔者未检索到该书公开出版的中文全译本。本项目译文仅用于专业学位硕士翻译实践、学术交流和论文案例分析，不作商业出版用途。",
        cites=True,
    )
    set_paragraph_text(doc, "从文本类型看", "从文本类型看，源文本并非传统意义上的科技说明文本，而是兼具信息功能、表达功能和论辩功能的跨学科学术文本。文本中既有无人机遥感、摄影测量、激光雷达、点云模型等技术内容，也有“行星性”“盖娅”“超客体”“后碳共同体”等理论概念；作者还通过艺术案例、隐喻和跨媒介互文推进生态批判。因此，本项目需要同时处理技术信息、理论概念、复杂句关系和批判性表达。")
    set_paragraph_text(doc, "项目执行过程中", "项目执行过程中，笔者参照《翻译服务 第1部分：笔译》（GB/T 19363.1—2022）[[CITE:22]]，结合学术翻译实践特点，形成“译前分析—术语准备—初译执行—双语自校—单语试读—终稿统校—术语归档”的基本流程。由于本项目由笔者个人承担，质量控制主要采取笔者自校、导师意见吸收和同伴试读反馈相结合的方式。", cites=True)
    set_paragraph_text(doc, "本项目的实施周期完整", "本项目的实施涵盖译前准备、译中实施与译后管理三个阶段。笔者搜集参考语料并整理约 150 条核心概念，形成双语术语表。在翻译阶段，笔者依托 Trados Studio 2022 完成初译、复杂句关系调整和首轮质量自检，随后结合导师意见、同伴试读反馈和双语自校结果，对术语、互文和论证关系进行多轮修订。全稿双语终审及项目归档日期为【待作者填写】。")
    set_paragraph_text(doc, "本项目的实施参照", "本项目参照《翻译服务 第1部分：笔译》（GB/T 19363.1—2022）的相关要求[[CITE:22]]，结合学术翻译的实际情况开展。鉴于源文本兼具技术知识、哲学讨论与生态批判，本项目在译前、译中和译后环节分别设置相应的准备、复核和质量控制步骤。", cites=True)
    set_paragraph_text(doc, "（1）项目分析与目标确立", "（1）项目分析与目标确立。项目启动后，笔者对约 17,000 词的源文本进行通读和问题标记，确认主要难点集中在跨学科理论概念、复杂句关系以及隐喻和互文表达。据此，项目以语义准确、术语一致、逻辑清楚和中文表达自然为基本目标。")
    set_paragraph_text(doc, "（3）技术资源与平行语料筹备", "（3）技术资源与平行语料筹备。笔者选用 Trados Studio 2022 作为主要计算机辅助翻译工具，建立项目文件、翻译记忆库和术语表；同时利用 MIT Press、JSTOR、CNKI、Google Scholar 等平台查阅无人机媒介研究、遥感技术、摄影史和环境人文学资料。处理人名、作品名和理论概念时，笔者对照相关文献确认来源和常见译法。例如，翻译哈伦·法罗基（Harun Farocki）相关段落时，笔者查阅其关于 operative images 的文章[[CITE:23]]，以判断“操作性图像”等表达的用法。", cites=True)
    set_paragraph_text(doc, "（4）术语表与格式规范制定", "（4）术语表与格式规范制定。笔者预先提取高频词、专有名词和跨学科词汇，编制约 150 条双语术语。每条记录包括英文原词、候选译名、暂定译法、上下文例句、参考来源和备注；同时统一专名加注、译注、引文和复杂句处理的基本格式。")
    set_paragraph_text(doc, "（5）项目进度规划与难点应对", "（5）项目进度规划与问题记录。笔者将项目划分为术语预处理、初译、技术自检、双语自校、单语试读、终稿修订和语料归档等阶段。对于译名冲突、复杂句关系不清和互文信息缺失等问题，笔者通过查阅平行文献、记录翻译日志、吸收导师意见和回到上下文复核等方式作出调整。")
    set_paragraph_text(doc, "初译阶段：", "初译阶段：笔者依托 Trados Studio 2022 逐段完成源文本翻译。一般叙述性内容优先保证语义准确和表达自然；理论概念密集的段落则先保留概念边界，再调整中文表达。遇到介词嵌套和复合长句时，先划分语义层次，再根据原文关系安排语序、标点和必要的连接语，同时调用术语表维持核心概念的一致。")
    set_paragraph_text(doc, "在质量反馈整理方面", "在质量反馈整理方面，笔者将修改意见归纳为术语与理论表达、复杂句与论证关系、隐喻修辞与语篇表达三类。对于试读中出现的逻辑卡顿，笔者回到英文原文核对因果、让步或转折关系，再决定是否调整语序或补出连接语。实质性修改在翻译日志中记录原译、终稿和修改理由，以便追溯。")
    set_paragraph_text(doc, "在处理争议性术语", "在处理争议性术语和理论表达时，笔者对 sensorium、planetarity、volumetric sensing 等核心概念比较候选译名，并查阅其理论来源和上下文用法。少数仅凭中文译名难以理解的概念，酌情保留原词或加入简短说明，以减少误读。")
    set_paragraph_text(doc, "最后，笔者从译后材料", "最后，笔者从译后材料中提取二十五个代表性案例，用于第三章分析。案例覆盖术语与理论表达、复杂句与论证关系、隐喻修辞与语篇表达三类问题。译后归档还包括术语表、翻译日志和翻译记忆等过程性材料。")

    # Chapter 4 and back matter.
    set_paragraph_text(doc, "第四章 总结与反思", "第四章 总结与反思", "Heading 1")
    set_paragraph_text(doc, "本章围绕《无人机感知与共同体》第三部分", "本章围绕《无人机感知与共同体》第三部分“地球”的英汉翻译实践，总结源文本的主要特征、翻译中出现的三类问题及其处理，并反思本项目仍需完善之处。")
    set_paragraph_text(doc, "首先，本报告对源语文本", "首先，源文本处于媒介研究、环境人文学与技术哲学的交叉领域，兼具知识传递与批判表达。作者在讨论无人机遥感、摄影测量和点云模型等技术内容时，也借助艺术作品、哲学概念和空间隐喻评价军事凝视、技术乐观主义和生态危机。因此，译文需要同时顾及概念准确、句内关系和语篇立场。")
    set_paragraph_text(doc, "其次，基于上述文本特点", "其次，本报告将实践中反复出现的问题整合为三类。第一类是跨学科术语与理论表达，重点在于 planetarity、sensorium、volumetric sensing、alterity 等概念的辨析和统一。第二类是复杂句与论证关系，重点在于识别主干、限定和逻辑推进。第三类是隐喻、修辞与语篇表达，重点在于保留图像和评价作用，并准确处理转折、对照和互文背景。")
    set_paragraph_text(doc, "针对这些问题", "针对上述问题，第三章以二十五个案例说明了三种相互配合的处理路径：通过查证来源和全文复核统一术语；通过调整语序、标点和连接方式梳理复杂句；通过保留关键形象、调整中文搭配和补充必要背景处理隐喻与语篇。各项分析均以终审译文为准，没有另行创造与附录不一致的“改译”。")
    set_paragraph_text(doc, "综合来看，本报告对引言", "综合来看，本报告对三个研究问题作出了回应。跨学科概念的译名需要由来源、语境和术语关系共同确定；复杂句处理需要先还原论证关系，再决定中文的句子边界和连接方式；隐喻与批判性表达的翻译则要判断其在段落中的作用，并在形象、语体和立场之间作出具体选择。")
    set_paragraph_text(doc, "通过本次翻译实践，笔者进一步认识到", "通过本次实践，笔者认识到，跨学科学术文本翻译需要在理解、查证、表达和复核之间反复往返。术语、句法和语篇不是彼此孤立的层面：一个术语可能决定整段论证的理解，一个连接方式也可能改变作者的立场强度。因此，案例分析应从原句中的真实问题出发，说明译文采取了什么处理以及该处理解决了什么问题，而不必为普通语言调整附加过多理论标签。")
    set_paragraph_text(doc, "与此同时，本次翻译实践仍有", "与此同时，本次实践仍有可完善之处。源文本涉及媒介研究、生态哲学、后殖民理论和无人机技术，部分概念仍需在更大范围的中文语料中持续核验；互文说明和关系明示也存在尺度问题，过少会造成理解障碍，过多则可能限制原文的开放性。后续审校仍应持续检查术语、专名、论证关系和译注的必要性。")
    set_paragraph_text(doc, "综上所述，本次《无人机感知与共同体》", "综上所述，本次《无人机感知与共同体》第三部分的英汉翻译实践形成了以术语查证与统一、句子重组与逻辑梳理、隐喻保留与语篇调整为核心的处理思路。这些结论限于本项目，但可以为媒介研究、环境人文学和技术哲学类文本的英汉翻译提供具体参照。")


def rebuild_references_and_backmatter(doc) -> None:
    ref_heading = find_paragraph(doc, exact="参考文献")
    ref_heading.style = "Heading 1"
    ack_heading = find_paragraph(doc, exact="致  谢")
    set_text(ack_heading, "致 谢", "Heading 1")
    body = ref_heading._p.getparent()
    elements = list(body)
    start = elements.index(ref_heading._p) + 1
    end = elements.index(ack_heading._p)
    ref_paras = []
    for el in elements[start:end]:
        if el.tag == qn("w:p"):
            ref_paras.append(next(p for p in doc.paragraphs if p._p is el))
    for i, entry in enumerate(REFERENCES):
        if i < len(ref_paras):
            set_text(ref_paras[i], entry, "References")
        else:
            p = doc.add_paragraph(entry, style="References")
            ack_heading._p.addprevious(p._p)
    for p in ref_paras[len(REFERENCES):]:
        remove_paragraph(p)

    # Acknowledgments: retain only a clear author-owned marker.
    appendix1 = find_paragraph(doc, exact="附录一 《无人机感知与共同体》原文与译文")
    current = ack_heading._p.getnext()
    ack_text = None
    while current is not None and current is not appendix1._p:
        nxt = current.getnext()
        if current.tag == qn("w:p"):
            if ack_text is None:
                ack_text = next(p for p in doc.paragraphs if p._p is current)
                set_text(ack_text, "【待作者填写】", "Normal")
                ack_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                current.getparent().remove(current)
        else:
            current.getparent().remove(current)
        current = nxt
    if ack_text is None:
        p = doc.add_paragraph("【待作者填写】", style="Normal")
        appendix1._p.addprevious(p._p)

    set_text(appendix1, "附录一 《无人机感知与共同体》原文与译文", "Heading 1")
    try:
        appendix2 = find_paragraph(doc, exact="附录二 主要术语对照表")
        remove_paragraph(appendix2)
    except KeyError:
        pass


def clear_container(container) -> None:
    for child in list(container._element):
        container._element.remove(child)
    p = OxmlElement("w:p")
    container._element.append(p)


def set_page_number_format(section, fmt: str | None, start: int | None = None) -> None:
    sect_pr = section._sectPr
    node = sect_pr.find(qn("w:pgNumType"))
    if fmt is None:
        if node is not None:
            sect_pr.remove(node)
        return
    if node is None:
        node = OxmlElement("w:pgNumType")
        sect_pr.append(node)
    node.set(qn("w:fmt"), fmt)
    if start is not None:
        node.set(qn("w:start"), str(start))


def add_page_field(paragraph, cached: str) -> None:
    append_simple_field(paragraph, "PAGE", cached)


def format_header_paragraph(paragraph, text: str) -> None:
    set_text(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_fixed_20(paragraph)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    for run in paragraph.runs:
        set_east_asia(run, "Songti SC", "Times New Roman", 10.5)


def configure_sections(doc) -> None:
    if len(doc.sections) != 3:
        raise RuntimeError(f"Expected 3 sections, got {len(doc.sections)}")
    doc.settings.odd_and_even_pages_header_footer = True
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.3)
        section.bottom_margin = Cm(3.3)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.8)
        section.gutter = Cm(0)
        section.header_distance = Cm(2.6)
        section.footer_distance = Cm(2.6)
        section.start_type = WD_SECTION.NEW_PAGE
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = False
        section.even_page_header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.even_page_footer.is_linked_to_previous = False

    # Cover and declaration section: no header/footer/page number.
    for container in [doc.sections[0].header, doc.sections[0].even_page_header, doc.sections[0].footer, doc.sections[0].even_page_footer]:
        clear_container(container)
    set_page_number_format(doc.sections[0], None)

    odd_header = "南京航空航天大学硕士学位论文"
    even_header = "《无人机感知与共同体》（第3部分）英汉翻译实践报告"
    for idx, cached in [(1, "I"), (2, "1")]:
        section = doc.sections[idx]
        clear_container(section.header)
        clear_container(section.even_page_header)
        format_header_paragraph(section.header.paragraphs[0], odd_header)
        format_header_paragraph(section.even_page_header.paragraphs[0], even_header)
        clear_container(section.footer)
        clear_container(section.even_page_footer)
        for footer in (section.footer, section.even_page_footer):
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.first_line_indent = Pt(0)
            set_fixed_20(p)
            add_page_field(p, cached)
            for run in p.runs:
                set_east_asia(run, "Songti SC", "Times New Roman", 10.5)
    set_page_number_format(doc.sections[1], "upperRoman", 1)
    set_page_number_format(doc.sections[2], "decimal", 1)


def set_update_fields(doc) -> None:
    # Page numbers and TOC caches are refreshed explicitly in Microsoft Word.
    # Forcing updates on open triggers a misleading external-link warning in
    # Word, so do not leave w:updateFields enabled in the delivery package.
    settings = doc.settings.element
    node = settings.find(qn("w:updateFields"))
    if node is not None:
        settings.remove(node)


def configure_cover_layout(doc) -> None:
    # Page-break anchors.
    find_paragraph(doc, starts="Nanjing University").paragraph_format.page_break_before = True
    find_paragraph(doc, exact="独创性声明").paragraph_format.page_break_before = True
    find_paragraph(doc, exact="ABSTRACT").paragraph_format.page_break_before = True
    find_paragraph(doc, exact="目录").paragraph_format.page_break_before = True
    # The body already begins with a section break; suppress the Heading 1
    # style's additional page break on this first paragraph.
    find_paragraph(doc, exact="第一章 引言").paragraph_format.page_break_before = False

    # Classification block.
    for p in doc.paragraphs[:2]:
        p.paragraph_format.first_line_indent = Pt(0)
        set_fixed_20(p)
        for run in p.runs:
            set_east_asia(run, "Songti SC", "Times New Roman", 10.5)

    cover_type = find_paragraph(doc, exact="专业学位硕士学位论文")
    cover_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_type.paragraph_format.first_line_indent = Pt(0)
    set_fixed_20(cover_type)
    for run in cover_type.runs:
        set_east_asia(run, "Songti SC", "Times New Roman", 36)

    for text in ["《无人机感知与共同体》", "（第3部分）英汉翻译实践报告"]:
        p = find_paragraph(doc, exact=text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        set_fixed_20(p)
        for run in p.runs:
            set_east_asia(run, "Heiti SC", "Times New Roman", 26)
            run.bold = True

    for text, east, size in [("南京航空航天大学", "FangSong_GB2312", 22), ("研究生院 外国语学院", "Songti SC", 12), ("【待作者填写】", "Songti SC", 12)]:
        matches = [p for p in doc.paragraphs if p.text.strip() == text]
        if not matches:
            continue
        p = matches[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        set_fixed_20(p)
        for run in p.runs:
            set_east_asia(run, east, "Times New Roman", size)

    if doc.tables:
        table = doc.tables[0]
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, top=60, bottom=60, start=60, end=60)
                p = cell.paragraphs[0]
                p.paragraph_format.first_line_indent = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.25
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                for run in p.runs:
                    set_east_asia(run, "Songti SC", "Times New Roman", 16)

    # English cover.
    for idx, p in enumerate(doc.paragraphs):
        if 15 <= idx <= 33:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            for run in p.runs:
                set_east_asia(run, "Songti SC", "Times New Roman", 12)
    title = find_paragraph(doc, starts="A Report on E-C Translation")
    for run in title.runs:
        set_east_asia(run, "Songti SC", "Times New Roman", 22)
        run.bold = True

    # Declaration page.
    for title_text in ("独创性声明", "使用授权声明"):
        p = find_paragraph(doc, exact=title_text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        set_fixed_20(p, before=10, after=10)
        for run in p.runs:
            set_east_asia(run, "Heiti SC", "Times New Roman", 14)
            run.bold = True
    for p in doc.paragraphs[34:42]:
        if p.text.strip() and p.text.strip() not in {"独创性声明", "使用授权声明"}:
            p.style = doc.styles["Normal"]
            set_fixed_20(p)

    # Front titles.
    for text in ("摘 要", "ABSTRACT"):
        p = find_paragraph(doc, exact=text)
        p.style = "FrontTocTitle"
    find_paragraph(doc, exact="目录").style = "FrontTitle"

    # Abstract text follows the same fixed 20-point line spacing as the body.
    abstract_start = find_paragraph(doc, exact="摘 要")
    toc_start = find_paragraph(doc, exact="目录")
    active = False
    for p in doc.paragraphs:
        if p._p is abstract_start._p:
            active = True
        if p._p is toc_start._p:
            break
        if not active or p.style.name in {"FrontTitle", "FrontTocTitle"}:
            continue
        p.style = doc.styles["Normal"]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_fixed_20(p)
        if p.text.strip().startswith(("关键词：", "Keywords:")):
            p.paragraph_format.first_line_indent = Pt(0)


def remove_old_toc_and_insert(doc) -> None:
    toc_title = find_paragraph(doc, exact="目录")
    body_start = find_paragraph(doc, exact="第一章 引言")
    # Remove everything after the TOC heading up to the section-break paragraph that precedes body_start.
    current = toc_title._p.getnext()
    while current is not None and current is not body_start._p:
        nxt = current.getnext()
        # Preserve the paragraph carrying the front-matter sectPr.
        ppr = current.find(qn("w:pPr")) if current.tag == qn("w:p") else None
        sect = ppr.find(qn("w:sectPr")) if ppr is not None else None
        if sect is not None:
            break
        current.getparent().remove(current)
        current = nxt
    if current is None:
        raise RuntimeError("Front-matter section break not found")
    section_break_anchor = current

    toc_paragraphs = []
    for title, level, scope in TOC_TITLES:
        cached = "I" if scope == "front" else "1"
        p = doc.add_paragraph(style=f"toc {min(level, 3)}")
        p.add_run(f"{title}\t{cached}")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        section_break_anchor.addprevious(p._p)
        toc_paragraphs.append(p)
    add_complex_field_start(toc_paragraphs[0], 'TOC \\o "1-3" \\h \\z \\t "FrontTocTitle,1"')
    add_complex_field_end(toc_paragraphs[-1])

    table_title = doc.add_paragraph("表目录", style="FrontTitle")
    # The one-page TOC fills the preceding page; allowing natural flow avoids
    # Word inserting an otherwise blank Roman-numeral page before the list.
    table_title.paragraph_format.page_break_before = False
    section_break_anchor.addprevious(table_title._p)
    table_entry = doc.add_paragraph("表3.1 源语文本的主要特征及其对翻译实践的影响\t1", style="TableListEntry")
    table_entry.paragraph_format.tab_stops.add_tab_stop(Cm(15.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    section_break_anchor.addprevious(table_entry._p)
    add_complex_field_start(table_entry, 'TOC \\h \\z \\c "表"')
    add_complex_field_end(table_entry)


def apply_body_formatting(doc) -> None:
    body_start = find_paragraph(doc, exact="第一章 引言")
    started = False
    for p in doc.paragraphs:
        if p._p is body_start._p:
            started = True
        if not started:
            continue
        if p.style.name not in {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Caption", "CaseLabel", "CaseText", "References"}:
            p.style = doc.styles["Normal"]
        if p.style.name == "Normal":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_fixed_20(p)
        for run in p.runs:
            if p.style.name in {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "CaseLabel"}:
                set_east_asia(run, "Heiti SC", "Times New Roman", {"Heading 1": 15, "Heading 2": 14, "Heading 3": 12, "Heading 4": 12, "CaseLabel": 10.5}[p.style.name])
                run.bold = True
            else:
                set_east_asia(run, "Songti SC", "Times New Roman", 10.5)


def normalize_all_fonts(doc) -> None:
    for p in doc.paragraphs:
        style = p.style.name
        if style in {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "CaseLabel"}:
            east = "Heiti SC"
            size = {"Heading 1": 15, "Heading 2": 14, "Heading 3": 12, "Heading 4": 12, "CaseLabel": 10.5}[style]
        elif style in {"FrontTitle", "FrontTocTitle"}:
            east, size = "Heiti SC", 14
        else:
            east, size = "Songti SC", 10.5
        for run in p.runs:
            bold = run.bold
            italic = run.italic
            superscript = run.font.superscript
            set_east_asia(run, east, "Times New Roman", size)
            run.bold = bold
            run.italic = italic
            run.font.superscript = superscript
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_east_asia(run, "Songti SC", "Times New Roman", run.font.size.pt if run.font.size else 10.5)


def strip_footnotes(docx_path: Path) -> None:
    fd, tmp_name = tempfile.mkstemp(suffix=".docx", prefix="mti_nofn_")
    os.close(fd)
    tmp_path = Path(tmp_name)
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = etree.fromstring(data)
                    for ref in root.xpath(".//w:footnoteReference", namespaces=NS):
                        run = ref.getparent()
                        parent = run.getparent()
                        parent.remove(run)
                    data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
                # Keep footnotes.xml and its relationship intact. Word requires
                # the original note records for package consistency even when
                # all visible body references have been removed.
                zout.writestr(item, data)
        os.replace(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def build(input_path: Path, output_path: Path) -> None:
    doc = Document(input_path)
    configure_styles(doc)
    # Remove the static TOC before heading edits so its entries cannot be
    # mistaken for body headings with the same text.
    remove_old_toc_and_insert(doc)
    update_front_and_body(doc)
    rebuild_chapter_three(doc)
    rebuild_references_and_backmatter(doc)
    configure_sections(doc)
    apply_body_formatting(doc)
    normalize_all_fonts(doc)
    configure_cover_layout(doc)
    set_update_fields(doc)
    doc.core_properties.title = "《无人机感知与共同体》（第3部分）英汉翻译实践报告"
    doc.core_properties.author = "薛扬"
    doc.core_properties.subject = "翻译硕士专业学位论文"
    doc.core_properties.keywords = "MTI, 学术文本翻译, 无人机感知"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    strip_footnotes(output_path)


def pdf_page_texts(pdf_path: Path) -> list[str]:
    reader = PdfReader(str(pdf_path))
    return [norm_text(page.extract_text() or "") for page in reader.pages]


def roman(number: int) -> str:
    vals = [(10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    result = ""
    for value, glyph in vals:
        while number >= value:
            result += glyph
            number -= value
    return result


def find_page(pages: list[str], needle: str, start: int = 0, end: int | None = None) -> int:
    target = norm_text(needle)
    if end is None:
        end = len(pages)
    for idx in range(start, end):
        if target in pages[idx]:
            return idx
    raise RuntimeError(f"Could not locate PDF page for {needle!r}")


def strip_complex_field(paragraph) -> None:
    for run in list(paragraph._p.findall(qn("w:r"))):
        has_field = run.find(qn("w:fldChar")) is not None or run.find(qn("w:instrText")) is not None
        if has_field:
            paragraph._p.remove(run)


def add_complex_field_start(paragraph, instruction: str) -> None:
    run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run.extend([begin, instr, separate])
    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, run)


def add_complex_field_end(paragraph) -> None:
    run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run.append(end)
    paragraph._p.append(run)


def finalize_toc(docx_path: Path, pdf_path: Path) -> dict:
    pages = pdf_page_texts(pdf_path)
    front_start = find_page(pages, "本翻译实践报告以丹麦学者凯瑟琳·毛雷尔")
    body_start = find_page(pages, "近年来，无人机技术在军事侦察")
    mapping: dict[str, str] = {}
    for title, level, scope in TOC_TITLES:
        if scope == "front":
            if title == "摘 要":
                page_idx = front_start
            elif title == "ABSTRACT":
                page_idx = find_page(pages, "This translation practice report examines", front_start, body_start)
            elif title == "目录":
                page_idx = find_page(pages, "目录", front_start, body_start)
            elif title == "表目录":
                page_idx = find_page(pages, "表目录", front_start, body_start)
            else:
                page_idx = find_page(pages, title, front_start, body_start)
            mapping[title] = roman(page_idx - front_start + 1)
        else:
            if title == "第一章 引言":
                page_idx = body_start
            elif title == "参考文献":
                page_idx = find_page(pages, "参考文献 ［1］Gregory", body_start)
            elif title == "致 谢":
                page_idx = find_page(pages, "致 谢 【待作者填写】", body_start)
            else:
                page_idx = find_page(pages, title, body_start)
            mapping[title] = str(page_idx - body_start + 1)
    table_page_idx = find_page(pages, "源语文本的主要特征及其对翻译实践的影响", body_start)
    table_page = str(table_page_idx - body_start + 1)

    doc = Document(docx_path)
    toc_title = find_paragraph(doc, exact="目录")
    table_title = find_paragraph(doc, exact="表目录")
    toc_paragraphs = []
    current = toc_title._p.getnext()
    while current is not None and current is not table_title._p:
        if current.tag == qn("w:p"):
            p = next(p for p in doc.paragraphs if p._p is current)
            if p.style.name.lower().startswith("toc "):
                toc_paragraphs.append(p)
        current = current.getnext()
    if len(toc_paragraphs) != len(TOC_TITLES):
        raise RuntimeError(f"TOC paragraph count mismatch: {len(toc_paragraphs)} != {len(TOC_TITLES)}")
    for p, (title, level, _) in zip(toc_paragraphs, TOC_TITLES):
        set_text(p, f"{title}\t{mapping[title]}", f"toc {min(level, 3)}")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        for run in p.runs:
            set_east_asia(run, "Songti SC", "Times New Roman", 10.5)
    add_complex_field_start(toc_paragraphs[0], 'TOC \\o "1-3" \\h \\z \\t "FrontTocTitle,1"')
    add_complex_field_end(toc_paragraphs[-1])

    entry = table_title._p.getnext()
    if entry is None or entry.tag != qn("w:p"):
        raise RuntimeError("Table-list entry missing")
    table_entry = next(p for p in doc.paragraphs if p._p is entry)
    set_text(table_entry, f"表3.1 源语文本的主要特征及其对翻译实践的影响\t{table_page}", "TableListEntry")
    table_entry.paragraph_format.tab_stops.add_tab_stop(Cm(15.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    for run in table_entry.runs:
        set_east_asia(run, "Songti SC", "Times New Roman", 10.5)
    add_complex_field_start(table_entry, 'TOC \\h \\z \\c "表"')
    add_complex_field_end(table_entry)
    set_update_fields(doc)
    doc.save(docx_path)
    return {"page_count": len(pages), "front_start": front_start + 1, "body_start": body_start + 1, "toc": mapping, "table_page": table_page}


def audit(docx_path: Path, appendix_path: Path) -> dict:
    doc = Document(docx_path)
    appendix = Document(appendix_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    source_corpus = " ".join(row.cells[0].text.replace("\n", " ") for row in appendix.tables[0].rows[1:])
    target_corpus = " ".join(row.cells[1].text.replace("\n", " ") for row in appendix.tables[0].rows[1:])
    cases_found = []
    for i, p in enumerate(doc.paragraphs):
        m = re.fullmatch(r"例\[(\d+)\]", p.text.strip())
        if not m:
            continue
        block = {"number": int(m.group(1))}
        for q in doc.paragraphs[i + 1 : i + 5]:
            for key, label in (("source", "原文："), ("simulated", "模拟初译："), ("final", "改译："), ("analysis", "分析：")):
                if q.text.startswith(label):
                    block[key] = q.text[len(label) :]
        cases_found.append(block)
    groups = {}
    current_group = None
    for p in doc.paragraphs:
        if p.text.startswith("3.3.1"):
            current_group = "3.3.1"
        elif p.text.startswith("3.3.2"):
            current_group = "3.3.2"
        elif p.text.startswith("3.3.3"):
            current_group = "3.3.3"
        elif p.text.startswith("3.4"):
            current_group = None
        elif current_group and re.fullmatch(r"例\[\d+\]", p.text.strip()):
            groups[current_group] = groups.get(current_group, 0) + 1
    case_checks = []
    for block, expected in zip(cases_found, CASES):
        case_checks.append(
            {
                "number": block.get("number"),
                "source_in_appendix": norm_text(block.get("source", "")) in norm_text(source_corpus),
                "final_in_appendix": norm_text(block.get("final", "")) in norm_text(target_corpus),
                "simulated_label": "simulated" in block,
            }
        )
    with zipfile.ZipFile(docx_path) as z:
        names = set(z.namelist())
        document_xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        settings_xml = z.read("word/settings.xml").decode("utf-8", errors="ignore")
    forbidden = ["谱系对位", "信息切分", "生成性学术术语", "多向度的感官统摄", "[XXXX", "二О二X", "翻译策略与解决方案"]
    return {
        "case_count": len(cases_found),
        "group_counts": groups,
        "source_truth_pass": sum(x["source_in_appendix"] for x in case_checks),
        "final_truth_pass": sum(x["final_in_appendix"] for x in case_checks),
        "case_checks": case_checks,
        "forbidden_hits": {term: text.count(term) for term in forbidden if term in text},
        "author_markers": text.count("【待作者填写】"),
        "heading_counts": {name: sum(1 for p in doc.paragraphs if p.style.name == name) for name in ("Heading 1", "Heading 2", "Heading 3")},
        "sections": [
            {
                "margins_cm": [round(s.top_margin.cm, 2), round(s.bottom_margin.cm, 2), round(s.left_margin.cm, 2), round(s.right_margin.cm, 2)],
                "page_cm": [round(s.page_width.cm, 2), round(s.page_height.cm, 2)],
            }
            for s in doc.sections
        ],
        "has_footnotes_part": "word/footnotes.xml" in names,
        "has_footnote_refs": "footnoteReference" in document_xml,
        "has_toc_field": "TOC \\o" in document_xml,
        "has_table_list_field": 'TOC \\h \\z \\c &quot;表&quot;' in document_xml or 'TOC \\h \\z \\c "表"' in document_xml,
        "has_update_fields": "updateFields" in settings_xml,
        "odd_even_headers": doc.settings.odd_and_even_pages_header_footer,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    sub = parser.add_subparsers(dest="command", required=True)
    p_build = sub.add_parser("build")
    p_build.add_argument("input", type=Path)
    p_build.add_argument("output", type=Path)
    p_toc = sub.add_parser("finalize-toc")
    p_toc.add_argument("docx", type=Path)
    p_toc.add_argument("pdf", type=Path)
    p_audit = sub.add_parser("audit")
    p_audit.add_argument("docx", type=Path)
    p_audit.add_argument("appendix", type=Path)
    args = parser.parse_args()
    if args.command == "build":
        build(args.input, args.output)
    elif args.command == "finalize-toc":
        print(json.dumps(finalize_toc(args.docx, args.pdf), ensure_ascii=False, indent=2))
    elif args.command == "audit":
        print(json.dumps(audit(args.docx, args.appendix), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
